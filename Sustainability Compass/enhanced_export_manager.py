#!/usr/bin/env python3
"""
Enhanced Export Manager with SDG Contribution Charts
Includes the specific SDG contribution chart requested by user
"""

import os
import json
from typing import Dict, List, Optional
from datetime import datetime
import pandas as pd
from export_manager import ReportExporter
from sdg_chart_generator import SDGContributionChart

# Try importing optional dependencies
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class EnhancedReportExporter(ReportExporter):
    """Enhanced report exporter with SDG contribution charts"""
    
    def __init__(self):
        super().__init__()
        self.chart_generator = SDGContributionChart()
    
    def export_enhanced_pdf_report(self, results: Dict, filename: str, language: str = 'en') -> bool:
        """Export PDF report with SDG contribution chart"""
        
        if not REPORTLAB_AVAILABLE:
            print("⚠️ ReportLab not available, using basic export")
            return super().export_pdf_report(results, filename, language)
        
        try:
            # Generate SDG contribution chart
            chart_path = None
            if results.get('sdg_mapping'):
                chart_path = self.chart_generator.create_sdg_contribution_chart(
                    results['sdg_mapping'], 
                    language, 
                    'temp_sdg_contribution_chart.png'
                )
            
            # Create PDF document
            doc = SimpleDocTemplate(filename, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            # Add title
            title = "تقرير تحليل الاستدامة الشامل" if language == 'ar' else "Comprehensive Sustainability Analysis Report"
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # Add analysis date
            date_text = f"تاريخ التحليل: {datetime.now().strftime('%Y-%m-%d')}" if language == 'ar' else f"Analysis Date: {datetime.now().strftime('%Y-%m-%d')}"
            story.append(Paragraph(date_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Add SDG contribution chart if available
            if chart_path and os.path.exists(chart_path):
                chart_title = "مساهمة التحليل الذكي في أهداف التنمية المستدامة" if language == 'ar' else "AI Sustainability Analysis Contribution to SDGs"
                story.append(Paragraph(chart_title, styles['Heading2']))
                story.append(Spacer(1, 12))
                
                # Add chart image
                try:
                    chart_img = Image(chart_path, width=6*inch, height=3*inch)
                    story.append(chart_img)
                    story.append(Spacer(1, 20))
                except Exception as e:
                    print(f"⚠️ Could not add chart to PDF: {e}")
            
            # Add executive summary
            self._add_executive_summary_to_pdf(story, results, styles, language)
            
            # Add ESG analysis
            self._add_esg_analysis_to_pdf(story, results, styles, language)
            
            # Add page break
            story.append(PageBreak())
            
            # Add SDG mapping details
            self._add_sdg_mapping_to_pdf(story, results, styles, language)
            
            # Add recommendations
            self._add_recommendations_to_pdf(story, results, styles, language)
            
            # Build PDF
            doc.build(story)
            
            # Clean up temporary chart file
            if chart_path and os.path.exists(chart_path):
                try:
                    os.remove(chart_path)
                except:
                    pass
            
            return True
            
        except Exception as e:
            print(f"❌ Enhanced PDF export failed: {str(e)}")
            # Fallback to basic export
            return super().export_pdf_report(results, filename, language)
    
    def export_enhanced_word_report(self, results: Dict, filename: str, language: str = 'en') -> bool:
        """Export Word report with SDG contribution chart"""
        
        if not DOCX_AVAILABLE:
            print("⚠️ python-docx not available, using basic export")
            return super().export_word_report(results, filename, language)
        
        try:
            # Generate SDG contribution chart
            chart_path = None
            if results.get('sdg_mapping'):
                chart_path = self.chart_generator.create_sdg_contribution_chart(
                    results['sdg_mapping'], 
                    language, 
                    'temp_sdg_contribution_chart.png'
                )
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = "تقرير تحليل الاستدامة الشامل" if language == 'ar' else "Comprehensive Sustainability Analysis Report"
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add analysis date
            date_text = f"تاريخ التحليل: {datetime.now().strftime('%Y-%m-%d')}" if language == 'ar' else f"Analysis Date: {datetime.now().strftime('%Y-%m-%d')}"
            doc.add_paragraph(date_text)
            
            # Add SDG contribution chart if available
            if chart_path and os.path.exists(chart_path):
                chart_title = "مساهمة التحليل الذكي في أهداف التنمية المستدامة" if language == 'ar' else "AI Sustainability Analysis Contribution to SDGs"
                doc.add_heading(chart_title, level=1)
                
                # Add chart image
                try:
                    doc.add_picture(chart_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"⚠️ Could not add chart to Word: {e}")
            
            # Add executive summary
            summary_title = "الملخص التنفيذي" if language == 'ar' else "Executive Summary"
            doc.add_heading(summary_title, level=1)
            exec_summary = results.get('executive_summary', 'No executive summary available.')
            doc.add_paragraph(exec_summary)
            
            # Add ESG Analysis
            esg_title = "تحليل الأداء البيئي والاجتماعي والحوكمة" if language == 'ar' else "ESG Performance Analysis"
            doc.add_heading(esg_title, level=1)
            
            esg_analysis = results.get('esg_analysis', {})
            for category, data in esg_analysis.items():
                if isinstance(data, dict):
                    category_title = f"{category.replace('_', ' ').title()}"
                    doc.add_heading(category_title, level=2)
                    
                    score = data.get('score', 'N/A')
                    doc.add_paragraph(f"Score: {score}/10")
                    
                    strengths = data.get('strengths', [])
                    if strengths:
                        doc.add_paragraph("Strengths:")
                        for strength in strengths:
                            doc.add_paragraph(f"• {strength}", style='List Bullet')
                    
                    weaknesses = data.get('weaknesses', [])
                    if weaknesses:
                        doc.add_paragraph("Areas for Improvement:")
                        for weakness in weaknesses:
                            doc.add_paragraph(f"• {weakness}", style='List Bullet')
            
            # Add SDG Mapping
            sdg_title = "ربط أهداف التنمية المستدامة" if language == 'ar' else "UN SDG Mapping"
            doc.add_heading(sdg_title, level=1)
            
            sdg_mapping = results.get('sdg_mapping', {})
            high_impact_sdgs = []
            medium_impact_sdgs = []
            
            for sdg_key, data in sdg_mapping.items():
                if sdg_key.startswith('sdg_') and isinstance(data, dict):
                    score = data.get('score', 0)
                    if score >= 7:
                        high_impact_sdgs.append((sdg_key, data))
                    elif score >= 4:
                        medium_impact_sdgs.append((sdg_key, data))
            
            if high_impact_sdgs:
                doc.add_heading("High Impact SDGs", level=2)
                for sdg_key, data in high_impact_sdgs:
                    sdg_num = sdg_key.split('_')[1]
                    name = data.get('name', f'SDG {sdg_num}')
                    score = data.get('score', 0)
                    doc.add_paragraph(f"SDG {sdg_num}: {name} (Score: {score}/10)")
                    
                    contributions = data.get('contributions', [])
                    if contributions:
                        for contrib in contributions:
                            doc.add_paragraph(f"• {contrib}", style='List Bullet')
            
            if medium_impact_sdgs:
                doc.add_heading("Medium Impact SDGs", level=2)
                for sdg_key, data in medium_impact_sdgs:
                    sdg_num = sdg_key.split('_')[1]
                    name = data.get('name', f'SDG {sdg_num}')
                    score = data.get('score', 0)
                    doc.add_paragraph(f"SDG {sdg_num}: {name} (Score: {score}/10)")
            
            # Add recommendations
            recommendations_title = "التوصيات" if language == 'ar' else "Recommendations"
            doc.add_heading(recommendations_title, level=1)
            
            recommendations = results.get('recommendations', [])
            for i, rec in enumerate(recommendations[:5], 1):
                doc.add_paragraph(f"{i}. {rec}")
            
            # Save document
            doc.save(filename)
            
            # Clean up temporary chart file
            if chart_path and os.path.exists(chart_path):
                try:
                    os.remove(chart_path)
                except:
                    pass
            
            return True
            
        except Exception as e:
            print(f"❌ Enhanced Word export failed: {str(e)}")
            # Fallback to basic export
            return super().export_word_report(results, filename, language)
    
    def _add_sdg_mapping_to_pdf(self, story: List, results: Dict, styles, language: str):
        """Add SDG mapping section to PDF with enhanced formatting"""
        
        section_title = "ربط أهداف التنمية المستدامة" if language == 'ar' else "UN SDG Mapping"
        story.append(Paragraph(section_title, styles['Heading1']))
        story.append(Spacer(1, 12))
        
        sdg_mapping = results.get('sdg_mapping', {})
        
        # Group SDGs by impact level
        high_impact = []
        medium_impact = []
        
        for sdg_key, data in sdg_mapping.items():
            if sdg_key.startswith('sdg_') and isinstance(data, dict):
                score = data.get('score', 0)
                if score >= 7:
                    high_impact.append((sdg_key, data))
                elif score >= 4:
                    medium_impact.append((sdg_key, data))
        
        # Add high impact SDGs
        if high_impact:
            story.append(Paragraph("High Impact SDGs (Score ≥ 7)", styles['Heading2']))
            story.append(Spacer(1, 8))
            
            for sdg_key, data in high_impact:
                sdg_num = sdg_key.split('_')[1]
                name = data.get('name', f'SDG {sdg_num}')
                score = data.get('score', 0)
                
                sdg_title = f"SDG {sdg_num}: {name} (Score: {score}/10)"
                story.append(Paragraph(sdg_title, styles['Heading3']))
                
                contributions = data.get('contributions', [])
                evidence = data.get('evidence', '')
                
                if contributions:
                    for contrib in contributions:
                        story.append(Paragraph(f"• {contrib}", styles['Normal']))
                
                if evidence:
                    story.append(Paragraph(f"Evidence: {evidence}", styles['Normal']))
                
                story.append(Spacer(1, 8))
        
        # Add medium impact SDGs
        if medium_impact:
            story.append(Paragraph("Medium Impact SDGs (Score 4-6)", styles['Heading2']))
            story.append(Spacer(1, 8))
            
            for sdg_key, data in medium_impact:
                sdg_num = sdg_key.split('_')[1]
                name = data.get('name', f'SDG {sdg_num}')
                score = data.get('score', 0)
                
                sdg_title = f"SDG {sdg_num}: {name} (Score: {score}/10)"
                story.append(Paragraph(sdg_title, styles['Heading3']))
                story.append(Spacer(1, 4))

def test_enhanced_export():
    """Test the enhanced export with SDG chart"""
    
    # Sample analysis results
    sample_results = {
        'executive_summary': 'This company shows strong performance in sustainability with significant contributions to multiple UN SDGs.',
        'esg_analysis': {
            'economic_financial_performance': {
                'score': 8.0,
                'strengths': ['Strong revenue growth', 'Transparent reporting'],
                'weaknesses': ['Limited financial diversity']
            },
            'environmental_performance': {
                'score': 7.5,
                'strengths': ['Renewable energy adoption', 'Carbon reduction initiatives'],
                'weaknesses': ['Water usage monitoring needed']
            },
            'social_performance': {
                'score': 8.2,
                'strengths': ['Employee wellness programs', 'Community engagement'],
                'weaknesses': ['Diversity reporting gaps']
            }
        },
        'sdg_mapping': {
            'sdg_4': {'score': 8.5, 'name': 'Quality Education', 'impact_level': 'High', 'contributions': ['Digital literacy programs']},
            'sdg_5': {'score': 8.0, 'name': 'Gender Equality', 'impact_level': 'High', 'contributions': ['Women leadership initiatives']},
            'sdg_7': {'score': 8.2, 'name': 'Affordable Clean Energy', 'impact_level': 'High', 'contributions': ['Solar power deployment']},
            'sdg_8': {'score': 7.8, 'name': 'Decent Work', 'impact_level': 'High', 'contributions': ['Job creation programs']},
            'sdg_9': {'score': 5.5, 'name': 'Innovation', 'impact_level': 'Medium', 'contributions': ['Technology development']},
            'sdg_10': {'score': 8.0, 'name': 'Reduced Inequalities', 'impact_level': 'High', 'contributions': ['Inclusive policies']},
            'sdg_12': {'score': 5.0, 'name': 'Responsible Consumption', 'impact_level': 'Medium', 'contributions': ['Waste reduction']},
            'sdg_13': {'score': 7.5, 'name': 'Climate Action', 'impact_level': 'High', 'contributions': ['Carbon neutrality goals']},
            'sdg_17': {'score': 5.8, 'name': 'Partnerships', 'impact_level': 'Medium', 'contributions': ['NGO collaborations']}
        },
        'recommendations': [
            'Enhance water conservation measures',
            'Improve diversity reporting transparency',
            'Expand renewable energy initiatives',
            'Strengthen supply chain sustainability',
            'Develop circular economy practices'
        ]
    }
    
    exporter = EnhancedReportExporter()
    
    # Test PDF export
    print("🧪 Testing enhanced PDF export...")
    pdf_success = exporter.export_enhanced_pdf_report(sample_results, 'test_enhanced_report.pdf', 'en')
    if pdf_success:
        print("✅ Enhanced PDF report created successfully!")
    else:
        print("❌ Enhanced PDF export failed")
    
    # Test Word export  
    print("🧪 Testing enhanced Word export...")
    word_success = exporter.export_enhanced_word_report(sample_results, 'test_enhanced_report.docx', 'en')
    if word_success:
        print("✅ Enhanced Word report created successfully!")
    else:
        print("❌ Enhanced Word export failed")

if __name__ == "__main__":
    test_enhanced_export() 