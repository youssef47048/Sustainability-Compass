# Export Manager for Sustainability Reports
# Import with fallbacks for optional dependencies
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
    ARABIC_SUPPORT = True
except ImportError:
    ARABIC_SUPPORT = False

from docx import Document
from docx.shared import Inches
import pandas as pd
import json
import os
from typing import Dict, List
from datetime import datetime
from config import SDG_GOALS, ESG_CATEGORIES

class ReportExporter:
    """
    Export sustainability analysis reports in multiple formats
    """
    
    def __init__(self):
        self.setup_fonts()
        
    def setup_fonts(self):
        """Setup fonts for Arabic text support"""
        try:
            # Try to register Arabic font (you may need to add a font file)
            # For now, we'll use built-in fonts with Unicode support
            pass
        except Exception:
            pass
    
    def export_pdf_report(self, analysis_results: Dict, output_path: str, language: str = 'en') -> bool:
        """
        Export comprehensive PDF report
        
        Args:
            analysis_results (Dict): Complete analysis results
            output_path (str): Output file path
            language (str): Report language ('en' or 'ar')
            
        Returns:
            bool: Success status
        """
        if not REPORTLAB_AVAILABLE:
            print("❌ PDF export requires reportlab package")
            print("Install with: pip install reportlab")
            return False
            
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4, 
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Build story
            story = []
            styles = getSampleStyleSheet()
            
            # Add custom styles for Arabic if needed
            if language == 'ar':
                arabic_style = ParagraphStyle(
                    'Arabic',
                    parent=styles['Normal'],
                    alignment=2,  # Right alignment for Arabic
                    fontName='Helvetica',
                    fontSize=12
                )
                styles.add(arabic_style)
            
            # Title
            title_text = "Sustainability Analysis Report" if language == 'en' else "تقرير تحليل الاستدامة"
            title = Paragraph(self._format_text(title_text, language), styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Executive Summary
            self._add_executive_summary_to_pdf(story, analysis_results, styles, language)
            story.append(PageBreak())
            
            # ESG Analysis
            self._add_esg_analysis_to_pdf(story, analysis_results, styles, language)
            story.append(PageBreak())
            
            # SDG Mapping
            self._add_sdg_mapping_to_pdf(story, analysis_results, styles, language)
            story.append(PageBreak())
            
            # Recommendations
            self._add_recommendations_to_pdf(story, analysis_results, styles, language)
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"PDF export failed: {str(e)}")
            return False
    
    def export_word_report(self, analysis_results: Dict, output_path: str, language: str = 'en') -> bool:
        """
        Export Word document report
        
        Args:
            analysis_results (Dict): Complete analysis results
            output_path (str): Output file path
            language (str): Report language ('en' or 'ar')
            
        Returns:
            bool: Success status
        """
        try:
            doc = Document()
            
            # Set document direction for Arabic
            if language == 'ar':
                doc.settings.rtl = True
            
            # Title
            title_text = "Sustainability Analysis Report" if language == 'en' else "تقرير تحليل الاستدامة"
            title = doc.add_heading(self._format_text(title_text, language), 0)
            
            # Executive Summary
            self._add_executive_summary_to_word(doc, analysis_results, language)
            doc.add_page_break()
            
            # ESG Analysis
            self._add_esg_analysis_to_word(doc, analysis_results, language)
            doc.add_page_break()
            
            # SDG Mapping
            self._add_sdg_mapping_to_word(doc, analysis_results, language)
            doc.add_page_break()
            
            # Recommendations
            self._add_recommendations_to_word(doc, analysis_results, language)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Word export failed: {str(e)}")
            return False
    
    def export_excel_report(self, analysis_results: Dict, output_path: str, language: str = 'en') -> bool:
        """
        Export Excel spreadsheet with detailed data
        
        Args:
            analysis_results (Dict): Complete analysis results
            output_path (str): Output file path
            language (str): Report language ('en' or 'ar')
            
        Returns:
            bool: Success status
        """
        try:
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # ESG Analysis Sheet
                self._create_esg_excel_sheet(analysis_results, writer, language)
                
                # SDG Mapping Sheet
                self._create_sdg_excel_sheet(analysis_results, writer, language)
                
                # Summary Sheet
                self._create_summary_excel_sheet(analysis_results, writer, language)
                
                # Recommendations Sheet
                self._create_recommendations_excel_sheet(analysis_results, writer, language)
            
            return True
            
        except Exception as e:
            print(f"Excel export failed: {str(e)}")
            return False
    
    def _format_text(self, text: str, language: str) -> str:
        """Format text for proper display based on language"""
        if language == 'ar' and ARABIC_SUPPORT:
            try:
                reshaped_text = arabic_reshaper.reshape(text)
                return get_display(reshaped_text)
            except:
                return text
        return text
    
    def _add_executive_summary_to_pdf(self, story: List, results: Dict, styles, language: str):
        """Add executive summary to PDF"""
        section_title = "Executive Summary" if language == 'en' else "الملخص التنفيذي"
        story.append(Paragraph(self._format_text(section_title, language), styles['Heading1']))
        story.append(Spacer(1, 12))
        
        exec_summary = results.get('executive_summary', 'No executive summary available.')
        
        # Debug: Add analysis metadata
        if 'analysis_metadata' in results:
            metadata = results['analysis_metadata']
            debug_info = f"Analysis Date: {metadata.get('analysis_date', 'Unknown')}\n"
            debug_info += f"Document Pages: {metadata.get('document_pages', 0)}\n"
            debug_info += f"Language: {metadata.get('language', 'Unknown')}\n"
            debug_info += f"API Calls Used: {metadata.get('api_calls_used', 0)}\n\n"
            
            debug_para = Paragraph(self._format_text(debug_info, language), styles['Normal'])
            story.append(debug_para)
        
        # Check if we have actual content or just error message
        if exec_summary and len(exec_summary) > 50:
            summary_para = Paragraph(self._format_text(exec_summary, language), styles['Normal'])
        else:
            # Add debug information if no real summary
            debug_text = f"Executive Summary Issue: {exec_summary}\n\n"
            debug_text += "Debug Information:\n"
            debug_text += f"- Results keys: {list(results.keys())}\n"
            debug_text += f"- Results type: {type(results)}\n"
            
            if results.get('error'):
                debug_text += f"- Error: {results.get('error_message', 'Unknown error')}\n"
            
            summary_para = Paragraph(self._format_text(debug_text, language), styles['Normal'])
        
        story.append(summary_para)
        story.append(Spacer(1, 12))
    
    def _add_esg_analysis_to_pdf(self, story: List, results: Dict, styles, language: str):
        """Add ESG analysis to PDF"""
        section_title = "ESG Performance Analysis" if language == 'en' else "تحليل الأداء البيئي والاجتماعي والحوكمة"
        story.append(Paragraph(self._format_text(section_title, language), styles['Heading1']))
        story.append(Spacer(1, 12))
        
        esg_analysis = results.get('esg_analysis', {})
        
        # Create ESG table
        esg_data = [['Category', 'Score', 'Strengths', 'Weaknesses'] if language == 'en' 
                   else ['الفئة', 'النتيجة', 'نقاط القوة', 'نقاط الضعف']]
        
        for category, data in esg_analysis.items():
            if isinstance(data, dict) and 'score' in data:
                category_name = category.title()
                score = str(data.get('score', 0))
                strengths = ', '.join(data.get('strengths', []))[:100]
                weaknesses = ', '.join(data.get('weaknesses', []))[:100]
                
                esg_data.append([
                    self._format_text(category_name, language),
                    score,
                    self._format_text(strengths, language),
                    self._format_text(weaknesses, language)
                ])
        
        if len(esg_data) > 1:
            esg_table = Table(esg_data, colWidths=[1.5*inch, 1*inch, 2*inch, 2*inch])
            esg_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(esg_table)
    
    def _add_sdg_mapping_to_pdf(self, story: List, results: Dict, styles, language: str):
        """Add SDG mapping to PDF"""
        section_title = "UN SDG Mapping" if language == 'en' else "ربط أهداف التنمية المستدامة"
        story.append(Paragraph(self._format_text(section_title, language), styles['Heading1']))
        story.append(Spacer(1, 12))
        
        sdg_mapping = results.get('sdg_mapping', {})
        
        # Create top SDGs table
        sdg_scores = []
        for sdg_key, data in sdg_mapping.items():
            if sdg_key.startswith('sdg_') and isinstance(data, dict):
                sdg_num = int(sdg_key.split('_')[1])
                score = data.get('score', 0)
                impact = data.get('impact_level', 'None')
                sdg_scores.append({
                    'sdg': f"SDG {sdg_num}: {SDG_GOALS.get(sdg_num, '')}",
                    'score': score,
                    'impact': impact
                })
        
        # Sort by score and get top 10
        sdg_scores.sort(key=lambda x: x['score'], reverse=True)
        top_sdgs = sdg_scores[:10]
        
        sdg_data = [['SDG', 'Score', 'Impact Level'] if language == 'en' 
                   else ['هدف التنمية المستدامة', 'النتيجة', 'مستوى التأثير']]
        
        for sdg in top_sdgs:
            sdg_data.append([
                self._format_text(sdg['sdg'][:60], language),  # Truncate long text
                str(sdg['score']),
                self._format_text(sdg['impact'], language)
            ])
        
        if len(sdg_data) > 1:
            sdg_table = Table(sdg_data, colWidths=[3*inch, 1*inch, 1.5*inch])
            sdg_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(sdg_table)
    
    def _add_recommendations_to_pdf(self, story: List, results: Dict, styles, language: str):
        """Add recommendations to PDF"""
        section_title = "Recommendations" if language == 'en' else "التوصيات"
        story.append(Paragraph(self._format_text(section_title, language), styles['Heading1']))
        story.append(Spacer(1, 12))
        
        recommendations = results.get('recommendations', [])
        
        for i, rec in enumerate(recommendations, 1):
            rec_text = f"{i}. {rec}"
            rec_para = Paragraph(self._format_text(rec_text, language), styles['Normal'])
            story.append(rec_para)
            story.append(Spacer(1, 6))
    
    def _add_executive_summary_to_word(self, doc, results: Dict, language: str):
        """Add executive summary to Word document"""
        section_title = "Executive Summary" if language == 'en' else "الملخص التنفيذي"
        doc.add_heading(self._format_text(section_title, language), level=1)
        
        exec_summary = results.get('executive_summary', 'No executive summary available.')
        para = doc.add_paragraph(self._format_text(exec_summary, language))
        
        if language == 'ar':
            para.alignment = 2  # Right alignment
    
    def _add_esg_analysis_to_word(self, doc, results: Dict, language: str):
        """Add ESG analysis to Word document"""
        section_title = "ESG Performance Analysis" if language == 'en' else "تحليل الأداء البيئي والاجتماعي والحوكمة"
        doc.add_heading(self._format_text(section_title, language), level=1)
        
        esg_analysis = results.get('esg_analysis', {})
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        headers = ['Category', 'Score', 'Strengths', 'Weaknesses'] if language == 'en' \
                 else ['الفئة', 'النتيجة', 'نقاط القوة', 'نقاط الضعف']
        
        for i, header in enumerate(headers):
            table.cell(0, i).text = self._format_text(header, language)
        
        # Data rows
        for category, data in esg_analysis.items():
            if isinstance(data, dict) and 'score' in data:
                row = table.add_row()
                row.cells[0].text = self._format_text(category.title(), language)
                row.cells[1].text = str(data.get('score', 0))
                row.cells[2].text = self._format_text(', '.join(data.get('strengths', [])), language)
                row.cells[3].text = self._format_text(', '.join(data.get('weaknesses', [])), language)
    
    def _add_sdg_mapping_to_word(self, doc, results: Dict, language: str):
        """Add SDG mapping to Word document"""
        section_title = "UN SDG Mapping" if language == 'en' else "ربط أهداف التنمية المستدامة"
        doc.add_heading(self._format_text(section_title, language), level=1)
        
        sdg_mapping = results.get('sdg_mapping', {})
        
        # Get top SDGs
        sdg_scores = []
        for sdg_key, data in sdg_mapping.items():
            if sdg_key.startswith('sdg_') and isinstance(data, dict):
                sdg_num = int(sdg_key.split('_')[1])
                score = data.get('score', 0)
                impact = data.get('impact_level', 'None')
                sdg_scores.append({
                    'sdg': f"SDG {sdg_num}: {SDG_GOALS.get(sdg_num, '')}",
                    'score': score,
                    'impact': impact
                })
        
        sdg_scores.sort(key=lambda x: x['score'], reverse=True)
        top_sdgs = sdg_scores[:15]  # Show more in Word
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 2'
        
        # Header row
        headers = ['SDG', 'Score', 'Impact Level'] if language == 'en' \
                 else ['هدف التنمية المستدامة', 'النتيجة', 'مستوى التأثير']
        
        for i, header in enumerate(headers):
            table.cell(0, i).text = self._format_text(header, language)
        
        # Data rows
        for sdg in top_sdgs:
            row = table.add_row()
            row.cells[0].text = self._format_text(sdg['sdg'], language)
            row.cells[1].text = str(sdg['score'])
            row.cells[2].text = self._format_text(sdg['impact'], language)
    
    def _add_recommendations_to_word(self, doc, results: Dict, language: str):
        """Add recommendations to Word document"""
        section_title = "Recommendations" if language == 'en' else "التوصيات"
        doc.add_heading(self._format_text(section_title, language), level=1)
        
        recommendations = results.get('recommendations', [])
        
        for i, rec in enumerate(recommendations, 1):
            para = doc.add_paragraph(f"{i}. ")
            run = para.add_run(self._format_text(rec, language))
            
            if language == 'ar':
                para.alignment = 2  # Right alignment
    
    def _create_esg_excel_sheet(self, results: Dict, writer, language: str):
        """Create ESG analysis Excel sheet"""
        esg_data = []
        esg_analysis = results.get('esg_analysis', {})
        
        for category, data in esg_analysis.items():
            if isinstance(data, dict):
                esg_data.append({
                    'Category': category.title(),
                    'Score': data.get('score', 0),
                    'Strengths': '; '.join(data.get('strengths', [])),
                    'Weaknesses': '; '.join(data.get('weaknesses', [])),
                    'Evidence': data.get('evidence', '')[:500]  # Truncate for Excel
                })
        
        if esg_data:
            df = pd.DataFrame(esg_data)
            df.to_excel(writer, sheet_name='ESG Analysis', index=False)
    
    def _create_sdg_excel_sheet(self, results: Dict, writer, language: str):
        """Create SDG mapping Excel sheet"""
        sdg_data = []
        sdg_mapping = results.get('sdg_mapping', {})
        
        for sdg_key, data in sdg_mapping.items():
            if sdg_key.startswith('sdg_') and isinstance(data, dict):
                sdg_num = int(sdg_key.split('_')[1])
                sdg_data.append({
                    'SDG Number': sdg_num,
                    'SDG Title': SDG_GOALS.get(sdg_num, ''),
                    'Score': data.get('score', 0),
                    'Impact Level': data.get('impact_level', 'None'),
                    'Contributions': '; '.join(data.get('contributions', [])),
                    'Improvement Areas': '; '.join(data.get('improvement_areas', [])),
                    'Evidence': data.get('evidence', '')[:300]
                })
        
        if sdg_data:
            df = pd.DataFrame(sdg_data)
            df = df.sort_values('Score', ascending=False)
            df.to_excel(writer, sheet_name='SDG Mapping', index=False)
    
    def _create_summary_excel_sheet(self, results: Dict, writer, language: str):
        """Create summary Excel sheet"""
        summary_data = {
            'Analysis Date': [datetime.now().strftime('%Y-%m-%d')],
            'Language': [language],
            'Document Pages': [results.get('analysis_metadata', {}).get('document_pages', 0)],
            'Has Tables': [results.get('analysis_metadata', {}).get('has_tables', False)]
        }
        
        # Add ESG averages
        esg_analysis = results.get('esg_analysis', {})
        esg_scores = [data.get('score', 0) for data in esg_analysis.values() if isinstance(data, dict)]
        if esg_scores:
            summary_data['Average ESG Score'] = [sum(esg_scores) / len(esg_scores)]
        
        # Add SDG averages
        sdg_mapping = results.get('sdg_mapping', {})
        sdg_scores = [data.get('score', 0) for data in sdg_mapping.values() 
                     if isinstance(data, dict) and data.get('score', 0) > 0]
        if sdg_scores:
            summary_data['Average SDG Score'] = [sum(sdg_scores) / len(sdg_scores)]
            summary_data['Active SDGs'] = [len(sdg_scores)]
        
        df = pd.DataFrame(summary_data)
        df.to_excel(writer, sheet_name='Summary', index=False)
    
    def _create_recommendations_excel_sheet(self, results: Dict, writer, language: str):
        """Create recommendations Excel sheet"""
        recommendations = results.get('recommendations', [])
        
        if recommendations:
            rec_data = [{'Recommendation': rec} for rec in recommendations]
            df = pd.DataFrame(rec_data)
            df.index = df.index + 1  # Start from 1
            df.to_excel(writer, sheet_name='Recommendations', index_label='#') 