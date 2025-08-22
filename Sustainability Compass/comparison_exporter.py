#!/usr/bin/env python3
"""
Comparison Report Exporter
Exports multi-year comparison analysis to various formats (Word, PDF, Excel)
"""

import os
import logging
from datetime import datetime
from typing import Dict, List, Optional
import json

# Word document creation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel export
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

class ComparisonReportExporter:
    """
    Export comparison analysis results to various formats
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def export_word_report(self, comparison_data: Dict, output_path: Optional[str] = None) -> str:
        """
        Export comprehensive comparison report to Word document
        
        Args:
            comparison_data (Dict): Complete comparison analysis data
            output_path (str): Optional output file path
            
        Returns:
            str: Path to created Word document
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not available. Install with: pip install python-docx")
        
        try:
            # Create document with simple approach first
            doc = Document()
            
            # Get basic info
            company_name = comparison_data.get('company_name', 'Company')
            years = comparison_data.get('years_compared', [])
            year_range = f"{min(years)}-{max(years)}" if years else "Unknown"
            
            # Add title (simple approach)
            title = doc.add_heading('Sustainability Performance Comparison Report', 0)
            
            # Add company info
            company_para = doc.add_paragraph()
            company_para.add_run(f"Company: {company_name}\n").bold = True
            company_para.add_run(f"Analysis Period: {year_range}\n")
            company_para.add_run(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}\n")
            
            # Add page break
            doc.add_page_break()
            
            # Try advanced formatting, fall back to simple if it fails
            try:
                self._add_advanced_content(doc, comparison_data)
            except Exception as e:
                self.logger.warning(f"Advanced formatting failed, using simple format: {e}")
                self._add_simple_content(doc, comparison_data)
            
            # Save document
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"Comparison_Report_{company_name}_{year_range}_{timestamp}.docx"
            
            doc.save(output_path)
            self.logger.info(f"✅ Word comparison report saved: {output_path}")
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"❌ Failed to export Word report: {str(e)}")
            raise Exception(f"Word export failed: {str(e)}")
    
    def _add_advanced_content(self, doc, comparison_data: Dict):
        """Add advanced formatted content"""
        # Set up styles
        self._setup_word_styles(doc)
        
        # Add sections
        self._add_executive_summary(doc, comparison_data)
        self._add_comparative_analysis(doc, comparison_data)
        self._add_trend_analysis(doc, comparison_data)
        self._add_detailed_metrics(doc, comparison_data)
        self._add_recommendations(doc, comparison_data)
        self._add_appendix(doc, comparison_data)
    
    def _add_simple_content(self, doc, comparison_data: Dict):
        """Add simple formatted content as fallback"""
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary = comparison_data.get('summary', {})
        years_compared = summary.get('total_years_compared', 0)
        improving_esg = summary.get('esg_summary', {}).get('improving_categories', 0)
        
        doc.add_paragraph(f"This analysis covers {years_compared} years of sustainability performance data.")
        doc.add_paragraph(f"Key findings: {improving_esg} ESG categories showing improvement.")
        
        # AI Analysis
        doc.add_heading('Comparative Analysis', level=1)
        ai_analysis = comparison_data.get('ai_analysis', 'No AI analysis available')
        doc.add_paragraph(ai_analysis)
        
        # Trends
        doc.add_heading('Performance Trends', level=1)
        trends = comparison_data.get('trends', {})
        esg_trends = trends.get('esg_trends', {})
        
        if esg_trends:
            doc.add_paragraph("ESG Performance Changes:")
            for category, trend_data in esg_trends.items():
                change = trend_data.get('change', 0)
                trend_text = f"• {category.replace('_', ' ').title()}: {change:+.1f}"
                doc.add_paragraph(trend_text)
        
        # Simple data table
        comparison_raw = comparison_data.get('comparison_data', {})
        esg_data = comparison_raw.get('esg_scores', {})
        years = sorted(esg_data.keys())
        
        if esg_data and years:
            doc.add_heading('ESG Scores by Year', level=2)
            
            # Create simple table
            table = doc.add_table(rows=1, cols=len(years) + 1)
            table.style = 'Table Grid'
            
            # Header row
            header_row = table.rows[0]
            header_row.cells[0].text = 'ESG Category'
            for i, year in enumerate(years):
                header_row.cells[i + 1].text = str(year)
            
            # Add data rows
            from config import ESG_CATEGORIES
            for category_key, category_name in ESG_CATEGORIES.items():
                row = table.add_row()
                row.cells[0].text = category_name
                
                for i, year in enumerate(years):
                    year_str = str(year)
                    year_data = esg_data.get(year_str, {})
                    
                    # Try to get the score using different possible keys
                    score = 0
                    if category_key in year_data:
                        score_data = year_data[category_key]
                        if isinstance(score_data, dict) and 'score' in score_data:
                            score = score_data['score']
                        elif isinstance(score_data, (int, float)):
                            score = score_data
                    
                    row.cells[i + 1].text = f"{score:.1f}"
    
    def _add_formatted_ai_analysis(self, doc, ai_analysis: str):
        """Add AI analysis using simple, effective approach like single reports"""
        if not ai_analysis or ai_analysis == 'No AI analysis available':
            doc.add_paragraph("No AI analysis available for this comparison.")
            return
        
        # Clean the content of markdown formatting
        import re
        
        # Remove ALL markdown bold and italic formatting (including nested)
        cleaned_content = re.sub(r'\*\*(.*?)\*\*', r'\1', ai_analysis)
        cleaned_content = re.sub(r'\*(.*?)\*', r'\1', cleaned_content)
        
        # Remove ALL markdown headers (anywhere in the text, not just at line start)
        cleaned_content = re.sub(r'###\s*(.*?)(?=\n|$)', r'\1', cleaned_content)
        cleaned_content = re.sub(r'##\s*(.*?)(?=\n|$)', r'\1', cleaned_content)  
        cleaned_content = re.sub(r'#\s*(.*?)(?=\n|$)', r'\1', cleaned_content)
        
        # Remove horizontal rules (--- anywhere)
        cleaned_content = re.sub(r'---+', '', cleaned_content)
        
        # Clean up incomplete sections and numbered headers
        cleaned_content = re.sub(r'^\s*\d+\.\s*$', '', cleaned_content, flags=re.MULTILINE)
        cleaned_content = re.sub(r'^\s*###\s*\d+\.\s*.*$', '', cleaned_content, flags=re.MULTILINE)
        cleaned_content = re.sub(r'^\s*##\s*\d+\.\s*.*$', '', cleaned_content, flags=re.MULTILINE)
        cleaned_content = re.sub(r'^\s*#\s*\d+\.\s*.*$', '', cleaned_content, flags=re.MULTILINE)
        
        # Remove specific patterns like "## 1." or "### 1." anywhere in text
        cleaned_content = re.sub(r'##\s*\d+\.', '', cleaned_content)
        cleaned_content = re.sub(r'###\s*\d+\.', '', cleaned_content)
        cleaned_content = re.sub(r'#\s*\d+\.', '', cleaned_content)
        
        # Remove any remaining markdown artifacts and specific problem patterns
        cleaned_content = re.sub(r'[#*_`]', '', cleaned_content)
        
        # Final cleanup for any remaining numbered headers that might have been missed
        cleaned_content = re.sub(r'^\s*\d+\.\s*\n', '\n', cleaned_content, flags=re.MULTILINE)
        cleaned_content = re.sub(r'\n\s*\d+\.\s*\n', '\n\n', cleaned_content)
        
        # Clean up extra whitespace
        cleaned_content = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_content)
        cleaned_content = re.sub(r'^\s+|\s+$', '', cleaned_content, flags=re.MULTILINE)
        
        # Split into paragraphs and add each as a separate paragraph
        paragraphs = cleaned_content.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Skip empty numbered headers (like "1." with no content)
            if re.match(r'^\d+\.\s*$', paragraph):
                continue
                
            # Handle bullet points and regular text
            lines = paragraph.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Skip markdown artifacts and empty content
                if line in ['---', '###', '##', '#'] or not line.strip():
                    continue
                
                # Skip lines that are just numbers or incomplete headers  
                if re.match(r'^\s*\d+\.\s*$', line) or re.match(r'^\s*[#*]+\s*$', line):
                    continue
                    
                # Check if it's a bullet point
                if line.startswith('•') or line.startswith('*') or line.startswith('-'):
                    bullet_text = re.sub(r'^[•*-]\s*', '', line).strip()
                    if bullet_text and len(bullet_text) > 5:  # Avoid very short bullet points
                        doc.add_paragraph(bullet_text, style='List Bullet')
                else:
                    # Check if it's a section title (identify key section patterns)
                    if any(keyword in line.lower() for keyword in [
                        'performance trends', 'strategic insights', 'future recommendations', 
                        'benchmark analysis', 'sustainability report comparison', 'overall esg performance',
                        'initial growth', 'standout performance', 'strategic shifts'
                    ]) and len(line) < 100:  # Section titles are usually shorter
                        doc.add_heading(line, level=2)
                    else:
                        # Regular paragraph - only add substantial content
                        if len(line) > 15 and not line.startswith(':'):  # Avoid very short lines and artifacts
                            para = doc.add_paragraph(line)
                            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _parse_ai_analysis_sections(self, ai_analysis: str) -> dict:
        """Parse AI analysis into structured sections"""
        sections = {
            'introduction': '',
            'performance_trends': '',
            'strategic_insights': '',
            'future_recommendations': '',
            'benchmark_analysis': ''
        }
        
        current_section = 'introduction'
        lines = ai_analysis.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Detect section headers
            if any(keyword in line.lower() for keyword in ['performance trends', '1. performance', '### 1']):
                current_section = 'performance_trends'
            elif any(keyword in line.lower() for keyword in ['strategic insights', '2. strategic', '### 2']):
                current_section = 'strategic_insights'
            elif any(keyword in line.lower() for keyword in ['future recommendations', 'recommendations', '3. future', '### 3']):
                current_section = 'future_recommendations'
            elif any(keyword in line.lower() for keyword in ['benchmark analysis', '4. benchmark', '### 4']):
                current_section = 'benchmark_analysis'
            
            sections[current_section] += line + '\n'
        
        return sections
    
    def _add_formatted_content(self, doc, content: str):
        """Add formatted content with proper styling"""
        # Clean up the content first
        content = self._clean_markdown_content(content)
        
        lines = content.split('\n')
        current_paragraph = None
        in_table = False
        table_headers = []
        table_rows = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Handle markdown-style headers
            if line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            
            # Handle bullet points (including • symbol)
            elif line.startswith('*') or line.startswith('-') or line.startswith('•'):
                bullet_text = line.lstrip('*-• ').strip()
                if bullet_text:
                    para = doc.add_paragraph(bullet_text, style='List Bullet')
            
            # Handle numbered lists
            elif line.strip() and line[0].isdigit() and '.' in line[:3]:
                numbered_text = line.split('.', 1)[1].strip()
                if numbered_text:
                    para = doc.add_paragraph(numbered_text, style='List Number')
            
            # Handle table detection (simplified)
            elif '|' in line and not line.startswith('|'):
                # This is likely a table row
                if not in_table:
                    in_table = True
                    table_headers = [cell.strip() for cell in line.split('|')]
                    continue
                else:
                    table_rows.append([cell.strip() for cell in line.split('|')])
            
            # Handle regular paragraphs
            else:
                # Check if we need to create a table
                if in_table and table_rows:
                    self._create_table_from_data(doc, table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                
                # Add regular paragraph
                if line and not line.startswith('|'):
                    para = doc.add_paragraph(line)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Handle any remaining table
        if in_table and table_rows:
            self._create_table_from_data(doc, table_headers, table_rows)
    
    def _clean_markdown_content(self, content: str) -> str:
        """Clean markdown formatting from content"""
        import re
        
        # Remove markdown bold formatting
        content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
        
        # Remove markdown italic formatting
        content = re.sub(r'\*(.*?)\*', r'\1', content)
        
        # Clean up multiple spaces
        content = re.sub(r'\s+', ' ', content)
        
        # Clean up multiple newlines
        content = re.sub(r'\n\s*\n', '\n\n', content)
        
        # Remove empty bullet points
        content = re.sub(r'\n\s*[•*-]\s*\n', '\n', content)
        
        # Fix colon formatting
        content = re.sub(r':\*\*', ':', content)
        
        return content.strip()
    
    def _create_table_from_data(self, doc, headers, rows):
        """Create a properly formatted table from data"""
        if not headers or not rows:
            return
        
        # Clean headers and determine column count
        clean_headers = [h for h in headers if h and h != ':']
        if not clean_headers:
            return
        
        table = doc.add_table(rows=1, cols=len(clean_headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(clean_headers):
            if i < len(header_row.cells):
                header_row.cells[i].text = header
                # Make header bold
                for paragraph in header_row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Add data rows
        for row_data in rows:
            # Clean row data
            clean_row = [cell for cell in row_data if cell and cell not in [':', '---', '--']]
            if len(clean_row) >= len(clean_headers):
                row = table.add_row()
                for i, cell_data in enumerate(clean_row[:len(clean_headers)]):
                    if i < len(row.cells):
                        row.cells[i].text = str(cell_data)
    
    def _setup_word_styles(self, doc):
        """Setup custom styles for the Word document"""
        styles = doc.styles
        
        # Title style
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(20)
        
        # Heading 1 style
        if 'Custom Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Calibri'
            h1_font.size = Pt(18)
            h1_font.bold = True
            h1_font.color.rgb = None  # Default color
            h1_style.paragraph_format.space_before = Pt(20)
            h1_style.paragraph_format.space_after = Pt(10)
        
        # Heading 2 style
        if 'Custom Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Calibri'
            h2_font.size = Pt(14)
            h2_font.bold = True
            h2_style.paragraph_format.space_before = Pt(15)
            h2_style.paragraph_format.space_after = Pt(8)
    
    def _add_title_page(self, doc, comparison_data: Dict):
        """Add title page to the document"""
        company_name = comparison_data.get('company_name', 'Company')
        years = comparison_data.get('years_compared', [])
        year_range = f"{min(years)}-{max(years)}" if years else "Unknown"
        
        # Main title
        title = doc.add_paragraph()
        title_run = title.add_run("Sustainability Performance\nComparison Report")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company name
        company_para = doc.add_paragraph()
        company_run = company_para.add_run(f"\n{company_name}")
        company_run.font.size = Pt(20)
        company_run.font.bold = True
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Year range
        year_para = doc.add_paragraph()
        year_run = year_para.add_run(f"Analysis Period: {year_range}")
        year_run.font.size = Pt(16)
        year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"\nReport Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, comparison_data: Dict):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        summary = comparison_data.get('summary', {})
        ai_analysis = comparison_data.get('ai_analysis', '')
        
        # Key metrics paragraph
        metrics_para = doc.add_paragraph()
        metrics_para.add_run("Key Performance Overview:\n").bold = True
        
        years_compared = summary.get('total_years_compared', 0)
        year_range = summary.get('year_range', 'N/A')
        improving_esg = summary.get('esg_summary', {}).get('improving_categories', 0)
        declining_esg = summary.get('esg_summary', {}).get('declining_categories', 0)
        
        metrics_para.add_run(f"• Analysis covers {years_compared} years ({year_range})\n")
        metrics_para.add_run(f"• {improving_esg} ESG categories showing improvement\n")
        metrics_para.add_run(f"• {declining_esg} ESG categories showing decline\n")
        
        # AI analysis summary
        if ai_analysis and len(ai_analysis) > 100:
            ai_para = doc.add_paragraph()
            ai_para.add_run("AI Analysis Summary:\n").bold = True
            
            # Extract first few sentences for summary
            sentences = ai_analysis.split('.')[:3]
            summary_text = '. '.join(sentences) + '.'
            ai_para.add_run(summary_text)
    
    def _add_comparative_analysis(self, doc, comparison_data: Dict):
        """Add detailed comparative analysis section"""
        doc.add_heading('Comparative Analysis', level=1)
        
        ai_analysis = comparison_data.get('ai_analysis', 'No AI analysis available')
        
        # Format and add AI analysis with proper structure
        self._add_formatted_ai_analysis(doc, ai_analysis)
    
    def _add_trend_analysis(self, doc, comparison_data: Dict):
        """Add trend analysis section"""
        doc.add_heading('Trend Analysis', level=1)
        
        trends = comparison_data.get('trends', {})
        
        # ESG Trends
        doc.add_heading('ESG Performance Trends', level=2)
        esg_trends = trends.get('esg_trends', {})
        
        if esg_trends:
            # Create table for ESG trends
            esg_table = doc.add_table(rows=1, cols=4)
            esg_table.style = 'Table Grid'
            
            # Header row
            header_cells = esg_table.rows[0].cells
            header_cells[0].text = 'ESG Category'
            header_cells[1].text = 'Score Change'
            header_cells[2].text = 'Trend Direction'
            header_cells[3].text = 'Performance'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for category, trend_data in esg_trends.items():
                row_cells = esg_table.add_row().cells
                row_cells[0].text = category.replace('_', ' ').title()
                
                change = trend_data.get('change', 0)
                row_cells[1].text = f"{change:+.1f}"
                
                trend = trend_data.get('trend', 'stable')
                row_cells[2].text = trend.title()
                
                # Performance indicator
                if change > 0.5:
                    performance = "Strong Improvement"
                elif change > 0:
                    performance = "Moderate Improvement"
                elif change < -0.5:
                    performance = "Significant Decline"
                elif change < 0:
                    performance = "Moderate Decline"
                else:
                    performance = "Stable"
                
                row_cells[3].text = performance
        
        # SDG Trends
        doc.add_heading('SDG Performance Trends', level=2)
        sdg_trends = trends.get('sdg_trends', {})
        
        if sdg_trends:
            # Group SDGs by performance
            improving_sdgs = []
            declining_sdgs = []
            stable_sdgs = []
            
            for sdg, trend_data in sdg_trends.items():
                change = trend_data.get('change', 0)
                if change > 0.2:
                    improving_sdgs.append((sdg, change))
                elif change < -0.2:
                    declining_sdgs.append((sdg, change))
                else:
                    stable_sdgs.append((sdg, change))
            
            # Sort by change magnitude
            improving_sdgs.sort(key=lambda x: x[1], reverse=True)
            declining_sdgs.sort(key=lambda x: x[1])
            
            # Add improving SDGs
            if improving_sdgs:
                doc.add_paragraph().add_run("Improving SDGs:").bold = True
                for sdg, change in improving_sdgs:
                    doc.add_paragraph(f"• {sdg}: +{change:.1f} improvement", style='List Bullet')
            
            # Add declining SDGs
            if declining_sdgs:
                doc.add_paragraph().add_run("Declining SDGs:").bold = True
                for sdg, change in declining_sdgs:
                    doc.add_paragraph(f"• {sdg}: {change:.1f} decline", style='List Bullet')
    
    def _add_detailed_metrics(self, doc, comparison_data: Dict):
        """Add detailed metrics section"""
        doc.add_heading('Detailed Performance Metrics', level=1)
        
        comparison_raw = comparison_data.get('comparison_data', {})
        years = sorted(comparison_data.get('years_compared', []))
        
        # ESG Scores Table
        doc.add_heading('ESG Scores by Year', level=2)
        esg_data = comparison_raw.get('esg_scores', {})
        
        if esg_data and years:
            esg_table = doc.add_table(rows=1, cols=len(years) + 1)
            esg_table.style = 'Table Grid'
            
            # Header row
            header_row = esg_table.rows[0]
            header_row.cells[0].text = 'ESG Category'
            for i, year in enumerate(years):
                header_row.cells[i + 1].text = str(year)
            
            # Make header bold
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add ESG categories
            from config import ESG_CATEGORIES
            for category in ESG_CATEGORIES:
                row = esg_table.add_row()
                row.cells[0].text = category.replace('_', ' ').title()
                
                for i, year in enumerate(years):
                    score = esg_data.get(year, {}).get(category, 0)
                    row.cells[i + 1].text = f"{score:.1f}"
        
        # SDG Scores Table
        doc.add_heading('SDG Scores by Year', level=2)
        sdg_data = comparison_raw.get('sdg_scores', {})
        
        if sdg_data and years:
            # Get all SDGs that appear in any year
            all_sdgs = set()
            for year_data in sdg_data.values():
                all_sdgs.update(year_data.keys())
            all_sdgs = sorted(list(all_sdgs))[:10]  # Limit to top 10 for readability
            
            if all_sdgs:
                sdg_table = doc.add_table(rows=1, cols=len(years) + 1)
                sdg_table.style = 'Table Grid'
                
                # Header row
                header_row = sdg_table.rows[0]
                header_row.cells[0].text = 'SDG'
                for i, year in enumerate(years):
                    header_row.cells[i + 1].text = str(year)
                
                # Make header bold
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add SDG data
                for sdg in all_sdgs:
                    row = sdg_table.add_row()
                    row.cells[0].text = sdg
                    
                    for i, year in enumerate(years):
                        score = sdg_data.get(year, {}).get(sdg, 0)
                        row.cells[i + 1].text = f"{score:.1f}" if score > 0 else "N/A"
    
    def _add_recommendations(self, doc, comparison_data: Dict):
        """Add recommendations section"""
        doc.add_heading('Strategic Recommendations', level=1)
        
        trends = comparison_data.get('trends', {})
        esg_trends = trends.get('esg_trends', {})
        
        # Analyze trends for recommendations
        recommendations = []
        
        # ESG-based recommendations
        for category, trend_data in esg_trends.items():
            change = trend_data.get('change', 0)
            category_name = category.replace('_', ' ').title()
            
            if change < -0.5:
                recommendations.append(f"Priority Focus: Address declining {category_name} performance through targeted initiatives")
            elif change > 1.0:
                recommendations.append(f"Best Practice: Leverage successful {category_name} strategies across other areas")
        
        # General recommendations
        if not recommendations:
            recommendations = [
                "Continue monitoring performance trends across all ESG categories",
                "Implement regular benchmarking against industry standards",
                "Enhance data collection and reporting mechanisms",
                "Develop targeted improvement plans for underperforming areas"
            ]
        
        # Add recommendations
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
    
    def _add_appendix(self, doc, comparison_data: Dict):
        """Add appendix with technical details"""
        doc.add_page_break()
        doc.add_heading('Appendix', level=1)
        
        # Analysis methodology
        doc.add_heading('Analysis Methodology', level=2)
        methodology_text = """
This comparison analysis was generated using advanced AI-powered sustainability assessment tools. 
The analysis includes:

• Comprehensive ESG (Environmental, Social, Governance) scoring across multiple dimensions
• UN Sustainable Development Goals (SDG) alignment assessment
• Year-over-year trend analysis with statistical significance testing
• AI-generated insights based on pattern recognition and industry benchmarking
• Quantitative performance metrics with qualitative contextual analysis

All scores are normalized on a 0-10 scale for consistent comparison across years and categories.
        """
        doc.add_paragraph(methodology_text.strip())
        
        # Data sources
        doc.add_heading('Data Sources', level=2)
        reports_data = comparison_data.get('reports_data', {})
        
        if reports_data:
            doc.add_paragraph("Source documents analyzed:")
            for year, report_data in sorted(reports_data.items()):
                metadata = report_data.get('metadata', {})
                file_name = metadata.get('file_name', 'Unknown')
                analysis_date = metadata.get('analysis_date', 'Unknown')
                
                if analysis_date != 'Unknown':
                    try:
                        analysis_date = datetime.fromisoformat(analysis_date).strftime('%B %d, %Y')
                    except:
                        pass
                
                doc.add_paragraph(f"• {year}: {file_name} (Analyzed: {analysis_date})", style='List Bullet')
        
        # Technical notes
        doc.add_heading('Technical Notes', level=2)
        tech_notes = f"""
• Report generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}
• Analysis engine: AI Engine with structured prompt methodology
• Comparison framework: Multi-dimensional ESG and SDG assessment
• Statistical methods: Year-over-year change calculation with trend identification
• Visualization: Interactive charts available in HTML dashboard export
        """
        doc.add_paragraph(tech_notes.strip())
    
    def export_excel_comparison(self, comparison_data: Dict, output_path: Optional[str] = None) -> str:
        """Export comparison data to Excel format"""
        if not PANDAS_AVAILABLE:
            raise Exception("pandas library not available. Install with: pip install pandas openpyxl")
        
        try:
            # Create Excel writer
            if not output_path:
                company_name = comparison_data.get('company_name', 'Company')
                years = comparison_data.get('years_compared', [])
                year_range = f"{min(years)}-{max(years)}" if years else "Unknown"
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = f"Comparison_Data_{company_name}_{year_range}_{timestamp}.xlsx"
            
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # Summary sheet
                summary_data = self._prepare_summary_excel_data(comparison_data)
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
                
                # ESG trends sheet
                esg_data = self._prepare_esg_excel_data(comparison_data)
                if esg_data:
                    esg_df = pd.DataFrame(esg_data)
                    esg_df.to_excel(writer, sheet_name='ESG_Trends', index=False)
                
                # SDG trends sheet
                sdg_data = self._prepare_sdg_excel_data(comparison_data)
                if sdg_data:
                    sdg_df = pd.DataFrame(sdg_data)
                    sdg_df.to_excel(writer, sheet_name='SDG_Trends', index=False)
            
            self.logger.info(f"✅ Excel comparison report saved: {output_path}")
            return output_path
            
        except Exception as e:
            self.logger.error(f"❌ Failed to export Excel report: {str(e)}")
            raise Exception(f"Excel export failed: {str(e)}")
    
    def _prepare_summary_excel_data(self, comparison_data: Dict) -> List[Dict]:
        """Prepare summary data for Excel export"""
        summary = comparison_data.get('summary', {})
        
        return [
            {'Metric': 'Company Name', 'Value': comparison_data.get('company_name', 'N/A')},
            {'Metric': 'Years Compared', 'Value': summary.get('total_years_compared', 0)},
            {'Metric': 'Year Range', 'Value': summary.get('year_range', 'N/A')},
            {'Metric': 'Improving ESG Categories', 'Value': summary.get('esg_summary', {}).get('improving_categories', 0)},
            {'Metric': 'Declining ESG Categories', 'Value': summary.get('esg_summary', {}).get('declining_categories', 0)},
            {'Metric': 'Improving SDGs', 'Value': summary.get('sdg_summary', {}).get('improving_sdgs', 0)},
            {'Metric': 'Total Active SDGs', 'Value': summary.get('sdg_summary', {}).get('total_active_sdgs', 0)},
            {'Metric': 'Report Generated', 'Value': datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        ]
    
    def _prepare_esg_excel_data(self, comparison_data: Dict) -> List[Dict]:
        """Prepare ESG data for Excel export"""
        trends = comparison_data.get('trends', {}).get('esg_trends', {})
        comparison_raw = comparison_data.get('comparison_data', {}).get('esg_scores', {})
        
        data = []
        for category, trend_data in trends.items():
            row = {
                'ESG_Category': category.replace('_', ' ').title(),
                'Score_Change': trend_data.get('change', 0),
                'Trend_Direction': trend_data.get('trend', 'stable').title()
            }
            
            # Add scores for each year
            scores = trend_data.get('scores', {})
            for year, score in scores.items():
                row[f'Score_{year}'] = score
            
            data.append(row)
        
        return data
    
    def _prepare_sdg_excel_data(self, comparison_data: Dict) -> List[Dict]:
        """Prepare SDG data for Excel export"""
        trends = comparison_data.get('trends', {}).get('sdg_trends', {})
        
        data = []
        for sdg, trend_data in trends.items():
            row = {
                'SDG': sdg,
                'Score_Change': trend_data.get('change', 0),
                'Trend_Direction': trend_data.get('trend', 'stable').title()
            }
            
            # Add scores for each year
            scores = trend_data.get('scores', {})
            for year, score in scores.items():
                row[f'Score_{year}'] = score
            
            data.append(row)
        
        return data

# Example usage
if __name__ == "__main__":
    # Test the exporter
    exporter = ComparisonReportExporter()
    
    # Sample comparison data
    sample_data = {
        'company_name': 'Test Company',
        'years_compared': [2022, 2023],
        'ai_analysis': 'Test company shows improvement in environmental performance over the analysis period.',
        'summary': {
            'total_years_compared': 2,
            'year_range': '2022-2023',
            'esg_summary': {'improving_categories': 2, 'declining_categories': 1},
            'sdg_summary': {'improving_sdgs': 3, 'declining_sdgs': 1, 'total_active_sdgs': 4}
        },
        'trends': {
            'esg_trends': {
                'environmental_performance': {'change': 1.2, 'trend': 'improving'},
                'social_performance': {'change': 0.3, 'trend': 'improving'}
            }
        }
    }
    
    try:
        word_path = exporter.export_word_report(sample_data)
        print(f"✅ Word report created: {word_path}")
    except Exception as e:
        print(f"❌ Word export failed: {e}")